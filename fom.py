from tkinter import *
from tkinter import ttk
from tkinter import filedialog as fd
from tkinter.messagebox import showinfo
import docx
from docx import Document
from docxtpl import DocxTemplate



window=Tk() #интерфейс
window.title("Добро пожаловать")
window.attributes("-toolwindow")


context = {}
developers = []

def open_file_text():   #функция открытия файла
    global context
    global developoers
    
    filetypes = (
        ('text files', '*.docx'),
        ('All files', '*.*')
    )

    filename = fd.askopenfilename(
        title='Open a file',
        initialdir='/',
        filetypes=filetypes)

    doc = docx.Document(filename)
    disc = doc.tables[1].rows[0].cells[3].text
    dire = doc.tables[1].rows[2].cells[2].text
    prof = doc.tables[1].rows[4].cells[4].text
    dean = doc.tables[0].rows[0].cells[1].text[:-1]
    i = 0
    while dean[i] != "_":
        i += 1
    while dean[i] != " ":
        i += 1
    dean = dean[i+1:]
    i = -1
    profes = []
    r = doc.tables[-1].rows[0].cells[4].text
    while ((r[1] == ".") and (r[3] == ".") and (r[4] == " ")) or ((r[-1] == ".") and (r[-3] == ".") and (r[-5] == " ")):
        developers.append(r)
        profes.append(doc.tables[-1].rows[0].cells[2].text)
        i -= 1
        if len(doc.tables[i].rows[0].cells) < 4:
            break
        r = doc.tables[i].rows[0].cells[4].text
    i = 1
    deve = developers[0] + ", " + profes[0]
    while (i < len(developers)):
        deve += "\n" + "\t" + "\t" + "  " + developers[1] + ", " + profes[1]
        i += 1
    context = {'discipline' : disc, 'direction' : dire, 'profile' : prof, 'dean' : dean, 'developer' : deve}
    shablon = Document("shablon.docx")
    r = shablon.tables[0].rows
    for i in range(1, len(doc.tables[3].rows)):
        row = shablon.tables[0].add_row()
        x = doc.tables[3].rows[i].cells[0].text
        j = 0
        while x[j] != " ":
            j += 1
        r[i].cells[1].text = x[:j]
        r[i].cells[2].text = x[j+1:]
        r[i].cells[3].text = doc.tables[3].rows[i].cells[1].text
        r[i].cells[4].text = doc.tables[3].rows[i].cells[2].text
    number = 1
    for i in range(2, len(shablon.tables[0].rows)):
        if r[i].cells[1].text == r[i-1].cells[1].text:
            r[i].cells[0].merge(r[i-1].cells[0])
            x = r[i].cells[1].text
            r[i].cells[1].merge(r[i-1].cells[1])
            r[i].cells[1].text = x
            x = r[i].cells[2].text
            r[i].cells[2].merge(r[i-1].cells[2])
            r[i].cells[2].text = x
        if r[i].cells[1].text != r[i-1].cells[1].text:
            r[i-1].cells[0].text = str(number)
            number += 1
    r[len(shablon.tables[0].rows)-1].cells[0].text = str(number)
    
    r = doc.tables[4].rows
    for i in range(len(r)):
        if len(r[i].cells[0].text) != 0:
            if (ord(r[i].cells[0].text[0]) > 48) and (ord(r[i].cells[0].text[0]) < 58):
                if len(r[i].cells[0].text) > 2:
                    if r[i].cells[0].text[1] != ".":
                        if (ord(r[i].cells[0].text[1]) < 49) or (ord(r[i].cells[0].text[1]) > 57):
                            row = shablon.tables[1].add_row()
                            shablon.tables[1].rows[-1].cells[0].text = r[i].cells[1].text
                    else:
                        if (ord(r[i].cells[0].text[2]) < 49) or (ord(r[i].cells[0].text[2]) > 57):
                            row = shablon.tables[1].add_row()
                            shablon.tables[1].rows[-1].cells[0].text = r[i].cells[1].text
                else:
                    if (ord(r[i].cells[0].text[0]) > 48) and (ord(r[i].cells[0].text[0]) < 58):
                        row = shablon.tables[1].add_row()
                        shablon.tables[1].rows[-1].cells[0].text = r[i].cells[1].text
    shablon.save("ФОМ.docx")
        
    showinfo(
        title='Selected File',
        message=filename
    )



main_lbl=Label(window, text="ВНИМАНИЕ! Открывайте окна для заполнения последовательно сверху вниз!", fg="red")
main_lbl.grid(column=0, row=0, columnspan=2)

open_text=Label(window, text='Загрузите файл РПД в расширении *.docx')
open_text.grid(column=0, row=1)

open_btn=Button(window, text='Выбрать файл', command=open_file_text, padx=15, pady=10)
open_btn.grid(column=1, row=1)

spisok=[]
def choose_os_page():
    choose=Toplevel()
    ch_os=Label(choose, text="Наименование оценочного средства")
    ch_os.grid(column=0, row=0)
    ch_oslist=Listbox(choose, selectmode=MULTIPLE, selectbackground="green", height=10, width=25)#список с выбором
    ch_oslist.grid(column=1, row=0)
    for i in ('Конспект лекций', 'Глоссарий по предмету', 'Тест', 'Устный опрос', 'Доклад/презентация', 'Реферат', 'Эссе', 'Контрольная работа', 'Практическое задание',
              'Решение задач', 'Лабораторная работа', 'Проект', 'Портфолио', 'Выставка', 'Деловая игра', 'Конференция', 'Олимпиада', 'Онлайн-курс'):             #заполнение списка
        ch_oslist.insert(END, i)
    def Final_list():
        global spisok
        spisok = []
        selectToSecond = ch_oslist.curselection()
        for i in selectToSecond:
            spisok.append(ch_oslist.get(i))

        global context
        shablon = DocxTemplate("ФОМ.docx")
        stroka=", ".join(spisok)
        cont = {'spisok' : stroka}
        context.update(cont)
        shablon.render(context)
        shablon.save("ФОМ.docx")
        
        choose.destroy()
    ch_save=Button(choose, text="Сохранить выбор", command=Final_list)
    ch_save.grid(column=1, row=1)
    choose.resizable(width=False, height=False)
    
index = 1
def table1_page():
    global index
    index = 1
    def next_row():
        global index
        
        spisok1=[]
        selectToSecond = table1_oslist.curselection()
        for i in selectToSecond:
            spisok1.append(table1_oslist.get(i))
        stroka=", ".join(spisok1)
        doc.tables[0].rows[index].cells[5].text = stroka
        doc.save("ФОМ.docx")
        
        index += 1
        if index == len(doc.tables[0].rows):
            table1_add.destroy()
            end_of_print=Label(table1, text="Все строки таблицы заполнены")
            end_of_print.grid(column=3, row=5)
            return
        table1_numb1.config(text=doc.tables[0].rows[index].cells[0].text)
        table1_index1.config(text=doc.tables[0].rows[index].cells[1].text)
        table1_comp1.config(text=doc.tables[0].rows[index].cells[2].text)
        table1_indexfull1.config(text=doc.tables[0].rows[index].cells[3].text)
        table1_result1.config(text=doc.tables[0].rows[index].cells[4].text)
        pass
    
    table1=Toplevel(window)
    table1_lbl=Label(table1, text='1. Заполнение паспорта фонда оценочных средств')#первая таблица
    table1_lbl.grid(column=3, row=2)
    table1_numb=Label(table1, text='№')
    table1_numb.grid(column=0, row=3)
    table1_index=Label(table1, text="Индекс компетенции")
    table1_index.grid(column=1, row=3)
    table1_comp=Label(table1, text='Содержание компетенции')
    table1_comp.grid(column=2,row=3)
    table1_indexfull=Label(table1, text="Индекс и содержание индикатора компетенции")
    table1_indexfull.grid(column=3, row=3)
    table1_result=Label(table1, text="Планируемые результаты обучения")
    table1_result.grid(column=4, row=3)
    table1_os=Label(table1, text="Наименование оценочного средства")
    table1_os.grid(column=5, row=3)

    table1_numb1=Label(table1, text='№', background="white")
    table1_numb1.grid(column=0, row=4)
    table1_index1=Label(table1, text="Индекс компетенции", background="white", width=20, wraplength=200)
    table1_index1.grid(column=1, row=4)
    table1_comp1=Label(table1, text='Содержание компетенции', background="white", width=30, wraplength=200)
    table1_comp1.grid(column=2,row=4)
    table1_indexfull1=Label(table1, text="Индекс и содержание индикатора компетенции",background="white", width=30, wraplength=200)
    table1_indexfull1.grid(column=3, row=4)
    table1_result1=Label(table1, text="Планируемые результаты обучения",background="white", width=30, wraplength=200)
    table1_result1.grid(column=4, row=4)
    table1_oslist=Listbox(table1, selectmode=MULTIPLE, selectbackground="green", height=5, width=25)#список с выбором

    doc = Document("ФОМ.docx")
    table1_numb1.config(text=doc.tables[0].rows[index].cells[0].text)
    table1_index1.config(text=doc.tables[0].rows[index].cells[1].text)
    table1_comp1.config(text=doc.tables[0].rows[index].cells[2].text)
    table1_indexfull1.config(text=doc.tables[0].rows[index].cells[3].text)
    table1_result1.config(text=doc.tables[0].rows[index].cells[4].text)
    
    for i in range(len(spisok)):
        table1_oslist.insert(END, spisok[i])
    table1_oslist.grid(column=5, row=4)
    table1_add=Button(table1, text="Сохранить строку", command=next_row)
    table1_add.grid(column=3, row=5)

    def table1_close():
        table1.destroy()
    
    table1_save=Button(table1, text="Завершить работу с таблицей", command=table1_close)
    table1_save.grid(column=3, row=6)
    table1.resizable(width=False, height=False)


def table2_page():
    global spisok
    global index
    index = 2
    def next_row():
        global index
        
        answers1=[]
        select1 = table2_code.curselection()
        for i in select1:
            answers1.append(table2_code.get(i))
        str1 = ", ".join(answers1)
        doc.tables[1].rows[index].cells[1].text = str1

        results = ""
        for ans in answers1:
            j = 1
            while ans != doc.tables[0].rows[j].cells[3].text:
                j += 1
            results += doc.tables[0].rows[j].cells[4].text + "\n" + "\n"
        doc.tables[1].rows[index].cells[2].text = results[:-2]
        
        doc.tables[1].rows[index].cells[3].text = table2_pok.get("1.0",END)

        doc.tables[1].rows[index].cells[4].text = table2_mark.get("1.0",END)
        
        answers2=[]
        select2 = table2_tk.curselection()
        for i in select2:
            answers2.append(table2_tk.get(i))
        str2 = ", ".join(answers2)
        doc.tables[1].rows[index].cells[5].text = str2

        answers3=[]
        select3 = table2_pa.curselection()
        for i in select3:
            answers3.append(table2_pa.get(i))
        str2 = ", ".join(answers3)
        doc.tables[1].rows[index].cells[6].text = str2
            
        doc.save("ФОМ.docx")
        
        index += 1
        if index == len(doc.tables[1].rows):
            table2_add.destroy()
            end_of_print2=Label(table2, text="Все строки таблицы заполнены")
            end_of_print2.grid(column=3, row=4)
            return
        table2_theme.config(text=doc.tables[1].rows[index].cells[0].text)
        pass
    
    table2=Toplevel(window)
    table2_lbl=Label(table2, text='2.1 Показатели и критерии оценивания компетенций')#вторая таблица
    table2_lbl.grid(column=0, row=0)
    table2_theme1=Label(table2, text='Тема или раздел дисциплины')
    table2_theme1.grid(column=0, row=1)
    table2_code1=Label(table2, text="Код индикатора компетенции")
    table2_code1.grid(column=1, row=1)
    table2_result1=Label(table2, text="Планируемый результат")
    table2_result1.grid(column=2,row=1)
    table2_pok1=Label(table2, text="Показатель")
    table2_pok1.grid(column=3, row=1)
    table2_mark1=Label(table2, text="Критерий оценивания")
    table2_mark1.grid(column=4, row=1)
    table2_nameos=Label(table2, text="Наименование ОС")
    table2_nameos.grid(column=5, row=1, columnspan=2)
    table2_tk1=Label(table2, text="ТК")
    table2_tk1.grid(column=5, row=2)
    table2_pa1=Label(table2, text="ПА")
    table2_pa1.grid(column=6, row=2)

    table2_theme=Label(table2, text='Тема или раздел дисциплины', background="white", wraplength=200)
    table2_theme.grid(column=0, row=3)
    table2_code=Listbox(table2, selectmode=MULTIPLE, selectbackground="green", height=5, width=25, exportselection=False)
    table2_code.grid(column=1, row=3)
    table2_result=Label(table2, text="<Планируемый результат>", background="white", wraplength=200)
    table2_result.grid(column=2,row=3)
    table2_pok=Text(table2, width=15, height=15)
    table2_pok.grid(column=3, row=3)
    table2_mark=Text(table2, width=15, height = 15)
    table2_mark.grid(column=4, row=3)
    table2_tk=Listbox(table2, selectmode=MULTIPLE, selectbackground="green", height=5, width=25, exportselection=False)
    table2_tk.grid(column=5, row=3)
    table2_pa= Listbox(table2, selectmode=MULTIPLE, selectbackground="green", height=5, width=25, exportselection=False)
    table2_pa.grid(column=6, row=3)
    table2.resizable(width=False, height=False)

    codes = []
    doc = Document("ФОМ.docx")
    table2_theme.config(text=doc.tables[1].rows[index].cells[0].text)
    
    for i in range(1, len(doc.tables[0].rows)):
        codes.append(doc.tables[0].rows[i].cells[3].text)
    
    for i in range(len(codes)):
        table2_code.insert(END, codes[i])

    for i in range(len(spisok)):
        table2_tk.insert(END, spisok[i])

    for i in range(len(spisok)):
        table2_pa.insert(END, spisok[i])
    
    table2_add=Button(table2, text="Сохранить строку", command=next_row)
    table2_add.grid(column=3, row=4)

    def table2_close():
        table2.destroy()
    
    table2_save=Button(table2, text="Завершить работу с таблицей", command=table2_close)
    table2_save.grid(column=3, row=5)
    
def table3_page():
    table3=Toplevel(window)
    table3_label=Label(table3, text="2.2 Критерии оценивания результатов обучения для текущего контроля успеваемости и промежуточной аттестации по дисциплине (модулю)")
    table3_label.grid(column=0, row=0, columnspan=3)
    table3_os=Label(table3, text="Оценочное средство")
    table3_os.grid(column=0, row=1)
    table3_krit=Label(table3, text="Критерии оценивания")
    table3_krit.grid(column=1, row=1)
    table3_mark=Label(table3, text="Шкала оценивания")
    table3_mark.grid(column=2, row=1)
    table3_oslist=Label(table3, text="Оценочное средство", background="white")
    table3_oslist.grid(column=0, row=2)
    table3_wrkrit=Text(table3, width=50)
    table3_wrkrit.grid(column=1, row=2)
    table3_wrmark=Text(table3, width=50)
    table3_wrmark.grid(column=2, row=2)
    table3_addos=Button(table3, text="Следующее оценочное средство", command="")
    table3_addos.grid(column=0, row=3)
    table3_add=Button(table3, text="Сохранить строку", command="")
    table3_add.grid(column=1, row=3)

    def table3_close():
        table3.destroy()
    table3_save=Button(table3, text="Завершить работу с таблицей", command=table3_close)
    table3_save.grid(column=1, row=4)
    table3.resizable(width=False, height=False)

def table4_page():
    table4=Toplevel(window)
    table4_label=Label(table4, text="")
    table4_label.grid(column=1, row=0)
    table4_numb=Label(table4, text="№")
    table4_numb.grid(column=0, row=1)
    table4_ans=Label(table4, text="Вопросы")
    table4_ans.grid(column=1, row=1)
    table4_comp=Label(table4, text="Код компетенции")
    table4_comp.grid(column=2, row=1)
    table4_numb=Label(table4, text="№", background="white")
    table4_numb.grid(column=0, row=2)
    table4_answer=Entry(table4)
    table4_answer.grid(column=1, row=2)
    table4_compet=Label(table4, text="Код компетенции", background="white")
    table4_compet.grid(column=2, row=2)
    table4_add=Button(table4, text="Сохранить строку")
    table4_add.grid(column=1, row=3)
    def table4_close():
        table4.destroy()
    table4_save=Button(table4, text="Завершить работу с таблицей", command=table4_close)
    table4_save.grid(column=1, row=4)
    table4.resizable(width=False, height=False)
    

choose_os=Label(window, text="Выбор оценочных средств")
choose_os.grid(column=0, row=3)
table1_label=Label(window, text="1. Паспорт фонда оценочных средств по дисциплине")
table1_label.grid(column=0, row=4)
table2_label=Label(window, text="2.1 Показатели и критерии оценивания компетенций")
table2_label.grid(column=0, row=5)
table3_label=Label(window, text="2.2 Критерии оценивания результатов обучения для текущего контроля успеваемости и промежуточной аттестации по дисциплине (модулю)")
table3_label.grid(column=0, row=6)
table4_label=Label(window, text="2.3 Вопросы к экзамену")
table4_label.grid(column=0, row=7)
choose_os_btn=Button(window, text="Открыть", command=choose_os_page)
choose_os_btn.grid(column=1, row=3)
table1_btn=Button(window, text="Открыть", command=table1_page)
table1_btn.grid(column=1, row=4)
table2_btn=Button(window, text="Открыть", command=table2_page)
table2_btn.grid(column=1, row=5)
table3_btn=Button(window, text="Открыть", command=table3_page)
table3_btn.grid(column=1, row=6)
table4_btn=Button(window, text="Открыть", command=table4_page)
table4_btn.grid(column=1, row=7)

def fom_close():
    str_="__________"
    doc = Document("ФОМ.docx")
    global developers
    for i in range(len(developers)):
        p = doc.add_paragraph(str_+developers[i]+"\n")
        
        
    p=doc.add_paragraph('«__________»')
    doc.save("ФОМ.docx")
    window.destroy()

fom_cl=Button(window, text="Завершить работу с ФОМ", command=fom_close)
fom_cl.grid(column=1, row=8)
window.resizable(width=False, height=False)

window.mainloop()
